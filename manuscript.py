"""
Manuscript export module for Story Forge.
Merges chapters into submission-ready document formats.

Supported export formats (selectable via checkbox):
- DOCX  — Primary manuscript (Word)
- TXT   — Plain text manuscript
- PDF   — Print-ready PDF
- ODT   — OpenDocument (LibreOffice)
- EPUB  — eBook distribution
- KPF   — Kindle Package Format (generated from EPUB via Kindle Previewer)
- KDP Proof PDF — Hardback/paperback proof for Amazon KDP
- KC    — Kindle Create project

Standard manuscript formatting:
- 12pt Times New Roman (or Courier)
- Double-spaced
- 1-inch margins
- Title page with author/title
- Chapter breaks with titles
- Page numbers
"""

import logging
import shutil
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak

from ebooklib import epub

from odf.opendocument import OpenDocumentText
from odf.style import Style, TextProperties, ParagraphProperties, PageLayoutProperties, PageLayout, MasterPage
from odf.text import P, H

from db import Book, Chapter, get_session
from sqlalchemy.orm import joinedload

logger = logging.getLogger(__name__)

# Output directory
MANUSCRIPT_DIR = Path(__file__).parent / "data" / "manuscripts"
MANUSCRIPT_DIR.mkdir(parents=True, exist_ok=True)

# All supported export formats
EXPORT_FORMATS = {
    "docx": {"label": "Word Document (.docx)", "direct": True},
    "txt": {"label": "Plain Text (.txt)", "direct": True},
    "pdf": {"label": "Print-Ready PDF (.pdf)", "direct": True},
    "odt": {"label": "OpenDocument (.odt)", "direct": True},
    "epub": {"label": "eBook (.epub)", "direct": True},
    "kdp_proof_pdf": {"label": "KDP Proof PDF (Hardback)", "direct": True},
    "kpf": {"label": "Kindle Package Format (.kpf)", "direct": False,
             "note": "Generate EPUB first, then convert via Kindle Previewer"},
    "kc": {"label": "Kindle Create Project", "direct": False,
            "note": "Import the DOCX into Kindle Create to produce this format"},
}


def get_book_with_chapters(book_id: int) -> Optional[Book]:
    """Load a book with all chapters ordered."""
    session = get_session()
    try:
        return session.query(Book).options(
            joinedload(Book.chapters)
        ).filter(Book.id == book_id).first()
    finally:
        session.close()


def _book_base_name(book: Book) -> str:
    """Generate a base filename like 'Title - Full Manuscript'."""
    title = book.title or "Untitled"
    return f"{title} - Full Manuscript"


def _safe_filename(name: str) -> str:
    """Make a filename safe for all OS."""
    return "".join(c if (c.isalnum() or c in ' -_().') else '_' for c in name).strip()


def _get_total_words(chapters: list) -> int:
    return sum(len(c.content.split()) if c.content else 0 for c in chapters)


# =============================================================================
# Export Package (multi-format with checkbox selection)
# =============================================================================


def export_package(
    book_id: int,
    formats: list[str],
    font_name: str = "Times New Roman",
    font_size: int = 12,
    double_spaced: bool = True,
    include_title_page: bool = True,
) -> dict:
    """
    Export a book in multiple formats as a package.

    Args:
        book_id: Book ID to export
        formats: List of format keys to export (e.g., ["docx", "pdf", "epub"])
        font_name: Font for manuscript formatting
        font_size: Font size in points
        double_spaced: Whether to double-space
        include_title_page: Whether to include title page

    Returns:
        dict with per-format results and package info
    """
    book = get_book_with_chapters(book_id)
    if not book:
        raise ValueError(f"Book not found: {book_id}")

    chapters = sorted(book.chapters, key=lambda c: c.order)
    if not chapters:
        raise ValueError(f"Book has no chapters: {book.title}")

    # Create a package directory for this export
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_title = _safe_filename(book.title)
    package_dir = MANUSCRIPT_DIR / f"{safe_title}_{timestamp}"
    package_dir.mkdir(parents=True, exist_ok=True)

    results = []
    notes = []
    export_kwargs = dict(
        book=book, chapters=chapters, output_dir=package_dir,
        font_name=font_name, font_size=font_size,
        double_spaced=double_spaced, include_title_page=include_title_page,
    )

    for fmt in formats:
        fmt_info = EXPORT_FORMATS.get(fmt)
        if not fmt_info:
            results.append({"format": fmt, "error": f"Unknown format: {fmt}"})
            continue

        if not fmt_info["direct"]:
            notes.append({"format": fmt, "note": fmt_info.get("note", "Requires external tool")})
            continue

        try:
            if fmt == "docx":
                result = _export_docx(**export_kwargs)
            elif fmt == "txt":
                result = _export_txt(**export_kwargs)
            elif fmt == "pdf":
                result = _export_pdf(**export_kwargs)
            elif fmt == "odt":
                result = _export_odt(**export_kwargs)
            elif fmt == "epub":
                result = _export_epub(**export_kwargs)
            elif fmt == "kdp_proof_pdf":
                result = _export_kdp_proof_pdf(**export_kwargs)
            else:
                result = {"format": fmt, "error": "Not implemented"}

            results.append(result)
        except Exception as e:
            logger.error(f"Failed to export {fmt}: {e}")
            results.append({"format": fmt, "error": str(e)})

    total_words = _get_total_words(chapters)

    return {
        "package_dir": str(package_dir),
        "book_title": book.title,
        "author": book.author or "Unknown",
        "word_count": total_words,
        "chapter_count": len(chapters),
        "page_estimate": max(1, total_words // 250),
        "exports": results,
        "notes": notes,
        "created_at": datetime.now().isoformat(),
    }


# =============================================================================
# DOCX Export
# =============================================================================


def _export_docx(book, chapters, output_dir, font_name, font_size, double_spaced, include_title_page) -> dict:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(font_size)
    if double_spaced:
        style.paragraph_format.line_spacing = 2.0
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    if include_title_page:
        _docx_title_page(doc, book, chapters, font_name, font_size)

    for i, chapter in enumerate(chapters):
        if i > 0 or include_title_page:
            doc.add_page_break()
        _docx_chapter(doc, chapter, font_name, font_size, double_spaced)

    base = _safe_filename(_book_base_name(book))
    filename = f"{base}.docx"
    path = output_dir / filename
    doc.save(str(path))

    return {"format": "docx", "filename": filename, "path": str(path), "size": path.stat().st_size}


def _docx_title_page(doc, book, chapters, font_name, font_size):
    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(book.title.upper())
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("by")
    run.font.name = font_name
    run.font.size = Pt(font_size)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(book.author or "Unknown Author")
    run.font.name = font_name
    run.font.size = Pt(font_size)

    for _ in range(4):
        doc.add_paragraph()

    total = _get_total_words(chapters)
    rounded = round(total / 1000) * 1000 if total > 1000 else total
    if rounded > 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Approximately {rounded:,} words")
        run.font.name = font_name
        run.font.size = Pt(font_size)


def _docx_chapter(doc, chapter, font_name, font_size, double_spaced):
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(chapter.title.upper())
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = True

    for _ in range(2):
        doc.add_paragraph()

    if chapter.content:
        for para_text in chapter.content.split('\n'):
            para_text = para_text.strip()
            if not para_text:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            if double_spaced:
                p.paragraph_format.line_spacing = 2.0
            run = p.add_run(para_text)
            run.font.name = font_name
            run.font.size = Pt(font_size)


# =============================================================================
# Single-file DOCX (backward compat)
# =============================================================================


def export_manuscript_docx(book_id, font_name="Times New Roman", font_size=12,
                           double_spaced=True, include_title_page=True) -> dict:
    book = get_book_with_chapters(book_id)
    if not book:
        raise ValueError(f"Book not found: {book_id}")
    chapters = sorted(book.chapters, key=lambda c: c.order)
    if not chapters:
        raise ValueError(f"Book has no chapters: {book.title}")

    result = _export_docx(book, chapters, MANUSCRIPT_DIR, font_name, font_size, double_spaced, include_title_page)
    total_words = _get_total_words(chapters)
    result.update({
        "word_count": total_words,
        "chapter_count": len(chapters),
        "page_estimate": max(1, total_words // 250),
        "created_at": datetime.now().isoformat(),
        "book_title": book.title,
        "author": book.author or "Unknown",
    })
    return result


# =============================================================================
# TXT Export
# =============================================================================


def _export_txt(book, chapters, output_dir, **_kwargs) -> dict:
    lines = []
    lines.append(book.title.upper())
    lines.append(f"by {book.author or 'Unknown Author'}")
    lines.append("")
    lines.append("---")
    lines.append("")

    for chapter in chapters:
        lines.append("")
        lines.append(chapter.title.upper())
        lines.append("")
        if chapter.content:
            for para in chapter.content.split('\n'):
                para = para.strip()
                if para:
                    lines.append(f"    {para}")
                    lines.append("")
        lines.append("")
        lines.append("* * *")
        lines.append("")

    base = _safe_filename(_book_base_name(book))
    filename = f"{base}.txt"
    path = output_dir / filename
    path.write_text("\n".join(lines), encoding="utf-8")

    return {"format": "txt", "filename": filename, "path": str(path), "size": path.stat().st_size}


def export_manuscript_txt(book_id) -> dict:
    book = get_book_with_chapters(book_id)
    if not book:
        raise ValueError(f"Book not found: {book_id}")
    chapters = sorted(book.chapters, key=lambda c: c.order)
    if not chapters:
        raise ValueError(f"Book has no chapters: {book.title}")

    result = _export_txt(book, chapters, MANUSCRIPT_DIR)
    total_words = _get_total_words(chapters)
    result.update({
        "word_count": total_words,
        "chapter_count": len(chapters),
        "page_estimate": max(1, total_words // 250),
        "created_at": datetime.now().isoformat(),
        "book_title": book.title,
        "author": book.author or "Unknown",
    })
    return result


# =============================================================================
# PDF Export
# =============================================================================


def _export_pdf(book, chapters, output_dir, font_name, font_size, double_spaced, include_title_page) -> dict:
    base = _safe_filename(_book_base_name(book))
    filename = f"{base}.pdf"
    path = output_dir / filename

    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        topMargin=1 * inch,
        bottomMargin=1 * inch,
        leftMargin=1 * inch,
        rightMargin=1 * inch,
    )

    styles = getSampleStyleSheet()

    # Map font name to reportlab font (Times New Roman -> Times-Roman)
    rl_font = "Times-Roman"
    if "courier" in font_name.lower():
        rl_font = "Courier"

    body_style = ParagraphStyle(
        'ManuscriptBody',
        parent=styles['Normal'],
        fontName=rl_font,
        fontSize=font_size,
        leading=font_size * (2.0 if double_spaced else 1.5),
        firstLineIndent=36,  # ~0.5 inch
        spaceBefore=0,
        spaceAfter=0,
    )

    title_style = ParagraphStyle(
        'ManuscriptTitle',
        parent=styles['Normal'],
        fontName=rl_font,
        fontSize=font_size,
        leading=font_size * 2,
        alignment=1,  # CENTER
    )

    chapter_heading_style = ParagraphStyle(
        'ChapterHeading',
        parent=styles['Normal'],
        fontName=rl_font,
        fontSize=font_size,
        leading=font_size * 2,
        alignment=1,
        spaceBefore=72,  # ~1 inch down
        spaceAfter=36,
    )

    story = []

    if include_title_page:
        story.append(Spacer(1, 3 * inch))
        story.append(Paragraph(f"<b>{book.title.upper()}</b>", title_style))
        story.append(Spacer(1, 24))
        story.append(Paragraph("by", title_style))
        story.append(Paragraph(book.author or "Unknown Author", title_style))
        story.append(Spacer(1, 2 * inch))
        total = _get_total_words(chapters)
        rounded = round(total / 1000) * 1000 if total > 1000 else total
        if rounded > 0:
            story.append(Paragraph(f"Approximately {rounded:,} words", title_style))
        story.append(PageBreak())

    for i, chapter in enumerate(chapters):
        if i > 0:
            story.append(PageBreak())

        story.append(Spacer(1, 1 * inch))
        story.append(Paragraph(f"<b>{chapter.title.upper()}</b>", chapter_heading_style))
        story.append(Spacer(1, 24))

        if chapter.content:
            for para_text in chapter.content.split('\n'):
                para_text = para_text.strip()
                if para_text:
                    story.append(Paragraph(para_text, body_style))

    doc.build(story)

    return {"format": "pdf", "filename": filename, "path": str(path), "size": path.stat().st_size}


# =============================================================================
# KDP Proof PDF (larger trim size, different margins for hardback)
# =============================================================================


def _export_kdp_proof_pdf(book, chapters, output_dir, font_name, font_size, double_spaced, include_title_page) -> dict:
    base = _safe_filename(f"kdp-proof-hardback")
    filename = f"{base}.pdf"
    path = output_dir / filename

    # KDP hardback trim: 6x9 inches is common
    page_width = 6 * inch
    page_height = 9 * inch

    doc = SimpleDocTemplate(
        str(path),
        pagesize=(page_width, page_height),
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        leftMargin=0.875 * inch,  # Slightly larger for gutter
        rightMargin=0.625 * inch,
    )

    styles = getSampleStyleSheet()
    rl_font = "Times-Roman"
    if "courier" in font_name.lower():
        rl_font = "Courier"

    body_size = 11  # Slightly smaller for book format

    body_style = ParagraphStyle(
        'KDPBody',
        parent=styles['Normal'],
        fontName=rl_font,
        fontSize=body_size,
        leading=body_size * (1.6 if double_spaced else 1.4),
        firstLineIndent=24,
        spaceBefore=0,
        spaceAfter=0,
    )

    title_style = ParagraphStyle(
        'KDPTitle',
        parent=styles['Normal'],
        fontName=rl_font,
        fontSize=16,
        leading=20,
        alignment=1,
    )

    chapter_style = ParagraphStyle(
        'KDPChapter',
        parent=styles['Normal'],
        fontName=rl_font,
        fontSize=14,
        leading=18,
        alignment=1,
        spaceBefore=54,
        spaceAfter=24,
    )

    story = []

    if include_title_page:
        story.append(Spacer(1, 2.5 * inch))
        story.append(Paragraph(f"<b>{book.title.upper()}</b>", title_style))
        story.append(Spacer(1, 18))
        story.append(Paragraph("by", title_style))
        story.append(Paragraph(book.author or "Unknown Author", title_style))
        story.append(PageBreak())

    for i, chapter in enumerate(chapters):
        if i > 0:
            story.append(PageBreak())

        story.append(Spacer(1, 0.75 * inch))
        story.append(Paragraph(f"<b>{chapter.title.upper()}</b>", chapter_style))
        story.append(Spacer(1, 18))

        if chapter.content:
            for para_text in chapter.content.split('\n'):
                para_text = para_text.strip()
                if para_text:
                    story.append(Paragraph(para_text, body_style))

    doc.build(story)

    return {"format": "kdp_proof_pdf", "filename": filename, "path": str(path), "size": path.stat().st_size}


# =============================================================================
# ODT Export (OpenDocument)
# =============================================================================


def _export_odt(book, chapters, output_dir, font_name, font_size, double_spaced, include_title_page) -> dict:
    doc = OpenDocumentText()

    # Page layout
    pl = PageLayout(name="PageLayout")
    pl.addElement(PageLayoutProperties(
        pagewidth="8.5in", pageheight="11in",
        margintop="1in", marginbottom="1in",
        marginleft="1in", marginright="1in",
    ))
    doc.automaticstyles.addElement(pl)

    mp = MasterPage(name="Standard", pagelayoutname=pl)
    doc.masterstyles.addElement(mp)

    # Text style
    body_style = Style(name="ManuscriptBody", family="paragraph")
    body_style.addElement(TextProperties(fontsize=f"{font_size}pt", fontname=font_name))
    if double_spaced:
        body_style.addElement(ParagraphProperties(linespacing="200%"))
    doc.styles.addElement(body_style)

    title_style = Style(name="ManuscriptTitle", family="paragraph")
    title_style.addElement(TextProperties(fontsize=f"{font_size}pt", fontname=font_name, fontweight="bold"))
    title_style.addElement(ParagraphProperties(textalign="center"))
    doc.styles.addElement(title_style)

    chapter_style = Style(name="ChapterHeading", family="paragraph")
    chapter_style.addElement(TextProperties(fontsize=f"{font_size}pt", fontname=font_name, fontweight="bold"))
    chapter_style.addElement(ParagraphProperties(textalign="center", margintop="1in"))
    doc.styles.addElement(chapter_style)

    if include_title_page:
        for _ in range(8):
            doc.text.addElement(P())

        p = P(stylename=title_style)
        p.addText(book.title.upper())
        doc.text.addElement(p)

        doc.text.addElement(P())

        p = P(stylename=title_style)
        p.addText("by")
        doc.text.addElement(p)

        p = P(stylename=title_style)
        p.addText(book.author or "Unknown Author")
        doc.text.addElement(p)

        # Page break after title
        pb_style = Style(name="PageBreak", family="paragraph")
        pb_style.addElement(ParagraphProperties(breakbefore="page"))
        doc.automaticstyles.addElement(pb_style)
        doc.text.addElement(P(stylename=pb_style))

    for i, chapter in enumerate(chapters):
        if i > 0:
            pb_style_n = Style(name=f"PB{i}", family="paragraph")
            pb_style_n.addElement(ParagraphProperties(breakbefore="page"))
            doc.automaticstyles.addElement(pb_style_n)
            doc.text.addElement(P(stylename=pb_style_n))

        p = P(stylename=chapter_style)
        p.addText(chapter.title.upper())
        doc.text.addElement(p)

        doc.text.addElement(P())

        if chapter.content:
            for para_text in chapter.content.split('\n'):
                para_text = para_text.strip()
                if para_text:
                    p = P(stylename=body_style)
                    p.addText(para_text)
                    doc.text.addElement(p)

    base = _safe_filename(_book_base_name(book))
    filename = f"{base}.odt"
    path = output_dir / filename
    doc.save(str(path))

    return {"format": "odt", "filename": filename, "path": str(path), "size": path.stat().st_size}


# =============================================================================
# EPUB Export
# =============================================================================


def _export_epub(book, chapters, output_dir, font_name, font_size, **_kwargs) -> dict:
    ebook = epub.EpubBook()

    # Metadata
    ebook.set_identifier(f"story-forge-{book.id}-{datetime.now().strftime('%Y%m%d')}")
    ebook.set_title(book.title)
    ebook.set_language('en')
    ebook.add_author(book.author or "Unknown Author")

    if book.description:
        ebook.add_metadata('DC', 'description', book.description)

    # CSS for manuscript styling
    css = f"""
    body {{ font-family: '{font_name}', serif; font-size: {font_size}pt; line-height: 1.8; margin: 1em; }}
    h1 {{ text-align: center; font-size: 1.2em; margin-top: 2em; margin-bottom: 1em; text-transform: uppercase; }}
    p {{ text-indent: 1.5em; margin: 0; padding: 0; }}
    .title-page {{ text-align: center; margin-top: 30%; }}
    .title-page h1 {{ font-size: 1.5em; }}
    .title-page .author {{ font-size: 1.1em; margin-top: 1em; }}
    """
    style = epub.EpubItem(uid="style", file_name="style/default.css", media_type="text/css", content=css.encode())
    ebook.add_item(style)

    spine = ['nav']
    toc = []

    # Title page
    title_html = f"""
    <html><head><link rel="stylesheet" href="../style/default.css" /></head>
    <body>
    <div class="title-page">
    <h1>{book.title}</h1>
    <p class="author">by {book.author or 'Unknown Author'}</p>
    </div>
    </body></html>
    """
    title_chapter = epub.EpubHtml(title='Title Page', file_name='title.xhtml', lang='en')
    title_chapter.content = title_html.encode()
    title_chapter.add_item(style)
    ebook.add_item(title_chapter)
    spine.append(title_chapter)

    # Chapters
    for chapter in chapters:
        content_html = ""
        if chapter.content:
            paragraphs = chapter.content.split('\n')
            for para in paragraphs:
                para = para.strip()
                if para:
                    # Escape HTML entities
                    para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    content_html += f"<p>{para}</p>\n"

        chapter_html = f"""
        <html><head><link rel="stylesheet" href="../style/default.css" /></head>
        <body>
        <h1>{chapter.title}</h1>
        {content_html}
        </body></html>
        """

        safe_ch_name = f"chapter_{chapter.order:03d}.xhtml"
        epub_chapter = epub.EpubHtml(title=chapter.title, file_name=safe_ch_name, lang='en')
        epub_chapter.content = chapter_html.encode()
        epub_chapter.add_item(style)
        ebook.add_item(epub_chapter)
        spine.append(epub_chapter)
        toc.append(epub_chapter)

    # Table of contents
    ebook.toc = toc
    ebook.add_item(epub.EpubNcx())
    ebook.add_item(epub.EpubNav())
    ebook.spine = spine

    base = _safe_filename(_book_base_name(book))
    filename = f"{base}.epub"
    path = output_dir / filename
    epub.write_epub(str(path), ebook)

    return {"format": "epub", "filename": filename, "path": str(path), "size": path.stat().st_size}


# =============================================================================
# Management
# =============================================================================


def list_manuscripts(book_id: Optional[int] = None) -> list[dict]:
    """List exported manuscripts and packages."""
    manuscripts = []

    for item in MANUSCRIPT_DIR.iterdir():
        if item.is_dir():
            # Package directory — list files inside
            files = [f.name for f in item.iterdir() if f.is_file()]
            manuscripts.append({
                "filename": item.name,
                "path": str(item),
                "is_package": True,
                "files": files,
                "format": "package",
                "size": sum(f.stat().st_size for f in item.iterdir() if f.is_file()),
                "modified_at": datetime.fromtimestamp(item.stat().st_mtime).isoformat(),
            })
        elif item.suffix in ('.docx', '.txt', '.pdf', '.odt', '.epub'):
            manuscripts.append({
                "filename": item.name,
                "path": str(item),
                "is_package": False,
                "format": item.suffix.lstrip('.'),
                "size": item.stat().st_size,
                "modified_at": datetime.fromtimestamp(item.stat().st_mtime).isoformat(),
            })

    manuscripts.sort(key=lambda x: x["modified_at"], reverse=True)
    return manuscripts


def delete_manuscript(filename: str) -> bool:
    """Delete a manuscript file or package directory."""
    path = MANUSCRIPT_DIR / filename
    if not path.exists() or path.parent != MANUSCRIPT_DIR:
        return False

    if path.is_dir():
        shutil.rmtree(str(path))
    else:
        path.unlink()
    return True


def get_available_formats() -> list[dict]:
    """Return all available export formats with metadata."""
    return [
        {"id": k, **v}
        for k, v in EXPORT_FORMATS.items()
    ]
